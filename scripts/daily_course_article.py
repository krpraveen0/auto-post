#!/usr/bin/env python3
"""Generate the next Medium course lesson as Markdown and DOCX."""

from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import sys
import textwrap
from pathlib import Path
from typing import Any

import requests

from course_agents import run_course_generation_with_agents


REQUIRED_SECTIONS = [
    "Learning Outcomes",
    "Worked Example",
    "Exercise",
    "Recap",
    "Next Lesson",
]


def dry_run_article(lesson: dict[str, str], previous_title: str | None, next_title: str | None) -> str:
    previous = previous_title or "Course index"
    next_lesson = next_title or "the next course series"
    sections = [
        f"# {lesson['title']}",
        "",
        f"Series navigation: Previous: {previous}. Course index: Agentic AI Engineering. Next: {next_lesson}.",
        "",
        "This is deterministic validation content generated without calling the OpenAI API. It exists to test Markdown output, DOCX writing, state handling, and render plumbing. Replace it with the model-generated article before publishing.",
        "",
        "## Learning Outcomes",
        "",
        f"1. Explain the lesson focus: {lesson['learns']}.",
        "2. Connect the idea to the broader course project.",
        "3. Complete a practical exercise with a clear expected output.",
        "",
        "## Worked Example",
        "",
        f"The lesson centers on {lesson['title']}. A useful Medium course installment should start from a concrete engineering situation, name the decision the reader needs to make, and then explain the concept with enough detail to help the reader act. In this validation article, the example is intentionally generic, but the structure mirrors the production requirement: context, decision, tradeoff, implementation detail, and reader takeaway.",
        "",
        "A strong lesson does not try to cover every related concept. It defines the boundary, shows one practical example, and gives the reader a small piece of work that moves the project forward. That pattern keeps the series useful for readers who are learning over multiple days instead of consuming one large reference guide.",
        "",
        "## Exercise",
        "",
        f"Exercise: {lesson['exercise']}.",
        "",
        "Expected output: a short artifact the reader can keep in their project repository, such as a decision table, configuration file, checklist, test case, or diagram note. The output should be concrete enough that the next lesson can build on it.",
        "",
        "## Recap",
        "",
        "This lesson should leave the reader with one durable mental model, one practical artifact, and one reason to continue. The course pipeline checks that every article includes learning outcomes, a worked example, an exercise, a recap, and a next-lesson bridge.",
        "",
        "## Next Lesson",
        "",
        f"Next, the series moves to {next_lesson}. The next lesson should reuse the artifact from this exercise rather than restarting from scratch.",
    ]
    filler = (
        "Course quality depends on continuity. Each lesson should preserve vocabulary from earlier lessons, avoid unexplained future terms, and make the project feel cumulative. "
        "The reader should not need to guess why the current lesson exists or how it connects to the final outcome. "
    )
    while len(re.findall(r"\b\w+\b", "\n".join(sections))) < 1250:
        sections.insert(-2, filler)
    return "\n".join(sections)


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "lesson"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def load_state(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"next_part": 1, "generated": [], "completion_issue_created": False}
    return json.loads(read_text(path))


def save_state(path: Path, state: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(state, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def parse_lessons(series_map: str) -> list[dict[str, str]]:
    lessons: list[dict[str, str]] = []
    in_sequence = False
    for line in series_map.splitlines():
        stripped = line.strip()
        if stripped == "## Lesson Sequence":
            in_sequence = True
            continue
        if in_sequence and stripped.startswith("## "):
            break
        if not in_sequence or not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if len(cells) < 6 or cells[0] in {"Part", "---:"} or set(cells[0]) <= {"-", ":"}:
            continue
        if not cells[0].isdigit():
            continue
        lessons.append(
            {
                "part": cells[0],
                "title": cells[1],
                "learns": cells[2],
                "prerequisite": cells[3],
                "exercise": cells[4],
                "visual": cells[5],
            }
        )
    return lessons


def extract_output_text(response: dict[str, Any]) -> str:
    if isinstance(response.get("output_text"), str):
        return response["output_text"].strip()
    chunks: list[str] = []
    for item in response.get("output", []):
        for content in item.get("content", []):
            if content.get("type") == "output_text" and content.get("text"):
                chunks.append(content["text"])
    return "\n".join(chunks).strip()


def call_openai(prompt: str, model: str) -> str:
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is required for scheduled article generation.")
    response = requests.post(
        "https://api.openai.com/v1/responses",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "input": prompt,
            "max_output_tokens": 7000,
        },
        timeout=180,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"OpenAI API error {response.status_code}: {response.text[:1000]}")
    text = extract_output_text(response.json())
    if not text:
        raise RuntimeError("OpenAI API returned no output text.")
    return text


def generate_article(prompt: str, model: str, backend: str, group_id: str | None) -> str:
    if backend == "responses":
        return call_openai(prompt, model)
    if backend == "agents":
        result = run_course_generation_with_agents(
            prompt,
            model=model,
            group_id=group_id,
            tracing_disabled=os.environ.get("OPENAI_AGENTS_DISABLE_TRACING") == "1",
        )
        return result.article
    raise ValueError(f"Unsupported GENERATION_BACKEND: {backend}. Use responses or agents.")


def quality_issues(article: str) -> list[str]:
    issues: list[str] = []
    word_count = len(re.findall(r"\b\w+\b", article))
    if word_count < 1200:
        issues.append(f"Article is too short for a useful course lesson: {word_count} words.")
    if word_count > 2600:
        issues.append(f"Article is too long for a focused Medium lesson: {word_count} words.")
    for section in REQUIRED_SECTIONS:
        if section.lower() not in article.lower():
            issues.append(f"Missing required section or phrase: {section}.")
    if "previous" not in article.lower() and "part 1" not in article.lower():
        issues.append("Missing previous/course navigation context.")
    if "next" not in article.lower():
        issues.append("Missing bridge to the next lesson.")
    return issues


def build_prompt(
    series_map: str,
    lesson_plan: str,
    lesson: dict[str, str],
    previous_title: str | None,
    next_title: str | None,
) -> str:
    return textwrap.dedent(
        f"""
        You are writing one installment in a high-quality Medium course series.

        Write a complete, publishable Medium lesson in Markdown.

        Series map:
        {series_map}

        Lesson plan:
        {lesson_plan}

        Current lesson:
        Part: {lesson['part']}
        Title: {lesson['title']}
        Reader learns: {lesson['learns']}
        Prerequisite: {lesson['prerequisite']}
        Exercise: {lesson['exercise']}
        Visual: {lesson['visual']}
        Previous lesson: {previous_title or "None"}
        Next lesson: {next_title or "None"}

        Requirements:
        - Target 1400-2200 words.
        - Teach one primary concept deeply.
        - Use clear, human, practical prose.
        - Include these sections: Learning Outcomes, Worked Example, Exercise, Recap, Next Lesson.
        - Include a short series navigation block near the top.
        - Include a concrete exercise with expected output.
        - Include visual guidance with caption and alt text.
        - Avoid hype, generic AI phrasing, and unsupported claims.
        - Mark factual claims that need sources with [SOURCE NEEDED: short note].
        - Do not invent URLs or citations.
        - Do not include tables unless a table is clearly better than prose.
        """
    ).strip()


def apply_docx_styles(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Generated Medium course lesson")


def add_markdown_line(document: Any, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("# "):
        document.add_heading(stripped[2:].strip(), level=0)
    elif stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=1)
    elif stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=2)
    elif re.match(r"^\d+\.\s+", stripped):
        document.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
    elif stripped.startswith("- "):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
    else:
        paragraph = document.add_paragraph()
        for idx, part in enumerate(re.split(r"(\*\*[^*]+\*\*)", stripped)):
            if not part:
                continue
            run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
            run.bold = part.startswith("**") and part.endswith("**")


def write_docx(article: str, output_path: Path, title: str) -> None:
    from docx import Document

    document = Document()
    apply_docx_styles(document)
    document.core_properties.title = title
    for line in article.splitlines():
        add_markdown_line(document, line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def append_github_output(values: dict[str, str]) -> None:
    output_path = os.environ.get("GITHUB_OUTPUT")
    if not output_path:
        return
    with open(output_path, "a", encoding="utf-8") as handle:
        for key, value in values.items():
            handle.write(f"{key}={value}\n")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--series-map", required=True, type=Path)
    parser.add_argument("--series-slug", required=True)
    parser.add_argument("--state-file", required=True, type=Path)
    parser.add_argument("--out-dir", required=True, type=Path)
    parser.add_argument("--model", default=os.environ.get("OPENAI_MODEL", "gpt-5"))
    parser.add_argument("--backend", default=os.environ.get("GENERATION_BACKEND", "responses"))
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    series_map = read_text(args.series_map)
    lessons = parse_lessons(series_map)
    if not lessons:
        raise RuntimeError(f"No lessons found in {args.series_map}.")

    state = load_state(args.state_file)
    next_part = int(state.get("next_part", 1))
    if next_part > len(lessons):
        append_github_output(
            {
                "generated": "false",
                "completed": "true",
                "completion_issue_needed": "true"
                if not state.get("completion_issue_created")
                else "false",
                "part": str(next_part),
            }
        )
        save_state(args.state_file, state)
        return 0

    lesson = lessons[next_part - 1]
    previous_title = lessons[next_part - 2]["title"] if next_part > 1 else None
    next_title = lessons[next_part]["title"] if next_part < len(lessons) else None

    lesson_plan_path = args.series_map.parent / f"{args.series_slug}-lesson-{next_part:02d}-plan.md"
    lesson_plan = read_text(lesson_plan_path) if lesson_plan_path.exists() else json.dumps(lesson, indent=2)

    prompt = build_prompt(series_map, lesson_plan, lesson, previous_title, next_title)
    if args.dry_run:
        article = dry_run_article(lesson, previous_title, next_title)
        issues = quality_issues(article)
    else:
        group_id = f"{args.series_slug}:part-{next_part:02d}"
        article = generate_article(prompt, args.model, args.backend, group_id)
        issues = quality_issues(article)
        if issues:
            repair_prompt = prompt + "\n\nFix these quality issues and return the complete revised lesson:\n- " + "\n- ".join(issues)
            article = generate_article(repair_prompt, args.model, args.backend, group_id)
            issues = quality_issues(article)
    if issues:
        raise RuntimeError("Generated article failed quality gate:\n- " + "\n- ".join(issues))

    today = dt.datetime.now(dt.timezone.utc).date().isoformat()
    lesson_slug = slugify(lesson["title"])
    base_name = f"{today}-part-{next_part:02d}-{lesson_slug}"
    md_path = args.out_dir / f"{base_name}.md"
    docx_path = args.out_dir / f"{base_name}.docx"
    manifest_path = args.out_dir / f"{base_name}.json"

    args.out_dir.mkdir(parents=True, exist_ok=True)
    md_path.write_text(article + "\n", encoding="utf-8")
    write_docx(article, docx_path, lesson["title"])
    manifest_path.write_text(
        json.dumps(
            {
                "series": args.series_slug,
                "part": next_part,
                "title": lesson["title"],
                "markdown": str(md_path),
                "docx": str(docx_path),
                "generated_at_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
                "model": args.model,
                "backend": args.backend,
            },
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )

    state.setdefault("generated", []).append(
        {
            "part": next_part,
            "title": lesson["title"],
            "docx": str(docx_path),
            "markdown": str(md_path),
            "generated_at_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
        }
    )
    state["next_part"] = next_part + 1
    save_state(args.state_file, state)
    append_github_output(
        {
            "generated": "true",
            "completed": "true" if next_part == len(lessons) else "false",
            "completion_issue_needed": "true"
            if next_part == len(lessons) and not state.get("completion_issue_created")
            else "false",
            "part": str(next_part),
            "title": lesson["title"],
            "docx_path": str(docx_path),
            "markdown_path": str(md_path),
            "manifest_path": str(manifest_path),
            "out_dir": str(args.out_dir),
        }
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise SystemExit(1)
